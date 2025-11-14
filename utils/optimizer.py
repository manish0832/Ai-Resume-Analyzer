import tempfile
import os
from docx import Document
import re

def generate_suggestions(resume_text, job_description, skills_analysis):
    """Generate improvement suggestions for the resume"""
    suggestions = []
    
    # ✅ Missing skills
    if skills_analysis['missing_skills']:
        missing = skills_analysis['missing_skills'][:5]
        suggestions.append({
            'type': 'skills',
            'title': 'Add Missing Technical Skills',
            'description': f"Consider adding these skills: {', '.join(missing)}",
            'priority': 'high'
        })
    
    # ✅ Keyword suggestions
    job_keywords = extract_important_keywords(job_description)
    resume_keywords = extract_important_keywords(resume_text)
    missing_keywords = list(set(job_keywords) - set(resume_keywords))

    if missing_keywords:
        suggestions.append({
            'type': 'keywords',
            'title': 'Include Relevant Keywords',
            'description': f"Add these keywords naturally: {', '.join(missing_keywords[:5])}",
            'priority': 'high'
        })
    
    # ✅ Format suggestions
    suggestions.extend(analyze_format_issues(resume_text))

    # ✅ Content suggestions
    suggestions.extend(analyze_content_issues(resume_text, job_description))

    return suggestions


# ✅ Improved keyword extractor
def extract_important_keywords(text):
    """Extract keywords from resume and JD"""
    words = re.findall(r'\b[A-Za-z]{3,}\b|\b[A-Z0-9]{2,}\b', text)

    stopwords = {
        'this','that','with','have','from','your','been','will','they','their','there',
        'into','about','such','very','more','than','then','them','over','also','only'
    }

    keywords = []
    for w in words:
        w = w.lower()
        if len(w) > 3 and w not in stopwords:
            keywords.append(w)

    return list(set(keywords))


# ✅ Improved format analysis
def analyze_format_issues(resume_text):
    suggestions = []
    text_lower = resume_text.lower()

    # ✅ Email check
    if not re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', resume_text):
        suggestions.append({
            'type': 'format',
            'title': 'Add Email Address',
            'description': 'Include a professional email address at the top of your resume.',
            'priority': 'high'
        })

    # ✅ India + Global phone formats
    if not re.search(r'(\+?\d{1,3}[- ]?)?\d{10}', resume_text):
        suggestions.append({
            'type': 'format',
            'title': 'Add Phone Number',
            'description': 'Include your phone number (e.g., +91 9876543210).',
            'priority': 'medium'
        })

    # ✅ Required sections
    required_sections = ['experience', 'education', 'skills']
    for section in required_sections:
        if section not in text_lower:
            suggestions.append({
                'type': 'format',
                'title': f'Missing {section.title()} Section',
                'description': f'Your resume should include a {section} section.',
                'priority': 'high'
            })

    # ✅ Length check
    wc = len(resume_text.split())
    if wc < 200:
        suggestions.append({
            'type': 'format',
            'title': 'Expand Resume Content',
            'description': 'Your resume is too short. Add more details and achievements.',
            'priority': 'medium'
        })
    elif wc > 1200:
        suggestions.append({
            'type': 'format',
            'title': 'Condense Resume Content',
            'description': 'Your resume appears too long. Keep only relevant information.',
            'priority': 'low'
        })

    return suggestions


# ✅ Improved content analysis
def analyze_content_issues(resume_text, job_description):
    suggestions = []
    text_lower = resume_text.lower()

    # ✅ Achievements check
    if not re.search(r'\d+%|\d+\s+(years?|months?)|\d{2,}', resume_text):
        suggestions.append({
            'type': 'content',
            'title': 'Add Quantifiable Achievements',
            'description': 'Include numbers or metrics to demonstrate your contributions.',
            'priority': 'high'
        })

    # ✅ Action verbs check
    action_verbs = [
        'managed','developed','created','implemented','improved','increased',
        'reduced','led','designed','built'
    ]
    found_verbs = [v for v in action_verbs if v in text_lower]

    if len(found_verbs) < 3:
        suggestions.append({
            'type': 'content',
            'title': 'Use Strong Action Verbs',
            'description': 'Start bullet points with words like "developed", "managed", "designed".',
            'priority': 'medium'
        })

    # ✅ Tech terminology check
    if any(term in job_description.lower() for term in ['developer', 'software', 'engineer']):
        tech_terms = ['api', 'database', 'algorithm', 'debug', 'framework', 'testing']
        found = [t for t in tech_terms if t in text_lower]

        if len(found) < 2:
            suggestions.append({
                'type': 'content',
                'title': 'Add Technical Terminology',
                'description': 'Include more technical terms relevant to engineering roles.',
                'priority': 'medium'
            })

    return suggestions


# ✅ SAFE DOCX GENERATOR (FINAL)
def create_optimized_resume(resume_text, suggestions):
    try:
        doc = Document()

        doc.add_heading("Optimized Resume", 0)
        doc.add_heading("Resume Content", level=1)

        # Remove unsupported characters
        safe_text = resume_text.encode("ascii", "ignore").decode()

        for paragraph in safe_text.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        doc.add_page_break()
        doc.add_heading("Optimization Suggestions", level=1)

        for i, sug in enumerate(suggestions, 1):
            doc.add_heading(f"{i}. {sug['title']}", level=2)
            doc.add_paragraph(f"Priority: {sug['priority'].upper()}")
            doc.add_paragraph(sug["description"])
            doc.add_paragraph("")

        # Save temp file
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()

        return tmp.name

    except Exception as e:
        raise Exception(f"Error creating optimized resume: {str(e)}")
